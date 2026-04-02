"""Word document report generator using python-docx."""
import io
import logging
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

COBALT = RGBColor(0x00, 0x2C, 0x77)
TEAL = RGBColor(0x00, 0xA8, 0xC7)


class DOCXGenerator:
    async def generate(self, context: Any) -> bytes:
        doc = Document()

        # Title
        title = doc.add_heading("Variance Analysis Report", level=0)
        title.runs[0].font.color.rgb = COBALT
        doc.add_paragraph(f"{context.period_id} | {context.view} vs {context.comparison_base}")

        # Financial Performance
        doc.add_heading("Financial Performance", level=1).runs[0].font.color.rgb = COBALT

        if context.summary_cards:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for i, h in enumerate(["Metric", "Actual", "Comparator", "Variance", "%"]):
                hdr[i].text = h

            for card in context.summary_cards:
                row = table.add_row().cells
                row[0].text = card.get("metric_name", "")
                row[1].text = f"${card.get('actual', 0):,.0f}"
                row[2].text = f"${card.get('comparator', 0):,.0f}"
                row[3].text = f"${card.get('variance_amount', 0):,.0f}"
                row[4].text = f"{card.get('variance_pct', 0):.1f}%"

        # Executive Summary Narrative (from narrative pyramid)
        doc.add_heading("Executive Summary", level=1).runs[0].font.color.rgb = COBALT
        exec_narr = context.executive_summary.get("full_narrative") or context.executive_summary.get("fullNarrative", "")
        if exec_narr:
            for para in exec_narr.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            # Fallback
            top_fav = next((v for v in context.variances if v.get("variance_amount", 0) > 0), None)
            if top_fav and top_fav.get("narrative_detail"):
                doc.add_paragraph(top_fav["narrative_detail"])

        # Section Narratives (Revenue, Costs, Profitability)
        if context.section_narratives:
            doc.add_heading("Section Analysis", level=1).runs[0].font.color.rgb = COBALT
            for section in context.section_narratives:
                s_name = section.get("section_name") or section.get("sectionName", "")
                s_narr = section.get("narrative", "")
                if s_name and s_narr:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{s_name}: ")
                    run.bold = True
                    p.add_run(s_narr)

        # Areas of Attention (upgraded to midlevel narratives)
        doc.add_heading("Areas of Attention", level=1).runs[0].font.color.rgb = COBALT
        unfavorable = sorted(
            [v for v in context.variances if v.get("variance_amount", 0) < 0],
            key=lambda v: abs(v.get("variance_amount", 0)), reverse=True,
        )[:5]
        for v in unfavorable:
            p = doc.add_paragraph()
            run = p.add_run(f"{v.get('account_name', '')}: ")
            run.bold = True
            # Use midlevel narrative (richer than oneliner)
            narr = v.get("narrative_midlevel") or v.get("narrativeMidlevel") or v.get("narrative_oneliner", "")
            p.add_run(narr)

        # Key Risks
        risks = context.executive_summary.get("key_risks") or context.executive_summary.get("keyRisks", [])
        if risks:
            doc.add_heading("Key Risks", level=1).runs[0].font.color.rgb = COBALT
            for risk in risks:
                r_text = risk.get("risk", str(risk)) if isinstance(risk, dict) else str(risk)
                doc.add_paragraph(f"- {r_text}")

        # Recommendations
        doc.add_heading("Outlook & Recommendations", level=1).runs[0].font.color.rgb = COBALT
        for i, alert in enumerate((context.trend_alerts or [])[:3], 1):
            doc.add_paragraph(f"{i}. {alert.get('description', '')} - {alert.get('projection', '')}")

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph("Marsh Vantage - Confidential - AI-Generated Report")
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x8B, 0x9A, 0xB5)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
